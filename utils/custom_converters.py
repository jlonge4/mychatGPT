from haystack import component
import docx
from typing import List, Union, Dict, Any, Optional
from haystack import Document
from pathlib import Path
import pandas as pd 


@component
class DocxToTextConverter:
    """
    A component to convert docx file to Document
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text
            doc = Document(content=text, meta=meta)
            documents.append(doc)
        return {"documents": documents}


@component
class ExcelToTextConverter:
    """
    A component to convert Excel (.xlsx) files to Document.
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            # Read the Excel file into a DataFrame
            df = pd.read_excel(file_path)
            # Convert the DataFrame to a string
            text = df.to_string(index=False)
            doc = Document(content=text, meta=meta)
            documents.append(doc)
        return {"documents": documents}
